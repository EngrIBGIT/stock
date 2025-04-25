import os
os.environ["CUDA_VISIBLE_DEVICES"] = "-1"  # Force CPU use for compatibility

import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import yfinance as yf
from tensorflow.keras.models import load_model
import xgboost as xgb
from sklearn.preprocessing import MinMaxScaler
from fpdf import FPDF
from docx import Document

# --- App Config ---
st.set_page_config(page_title="Stock Comparison App", layout="wide")
st.title("📈 Stock Comparison and Prediction App")

# --- Styles ---
st.markdown("""
<style>
.main {background-color: #e0f7fa;}
.sidebar .sidebar-content {background-color: #d8f3dc;}
.css-1q8dd3e {background-color: #fff3cd;}
.marquee {
    font-size: 18px;
    color: #0c0c0c;
    padding: 10px;
    background-color: #ffffe0;
    overflow: hidden;
    white-space: nowrap;
    animation: marquee 20s linear infinite;
}
@keyframes marquee {
    0% {transform: translateX(100%);}
    100% {transform: translateX(-100%);}
}
</style>
""", unsafe_allow_html=True)

# --- Load Models Automatically ---
@st.cache_resource
def load_lstm_model(path):
    return load_model(path, compile=False)

@st.cache_resource
def load_xgb_model(path):
    model = xgb.Booster()
    model.load_model(path)
    return model

# Load models from local files
lstm_models = {
    "PEP": load_lstm_model("models/pep_lstm_model.h5") if os.path.exists("models/pep_lstm_model.h5") else None,
    "KO": load_lstm_model("models/ko_lstm_model.h5") if os.path.exists("models/ko_lstm_model.h5") else None
}

xgb_models = {
    "PEP": load_xgb_model("models/xgb_pep_model.xgb") if os.path.exists("models/xgb_pep_model.xgb") else None,
    "KO": load_xgb_model("models/xgb_ko_model.xgb") if os.path.exists("models/xgb_ko_model.xgb") else None
}

# --- Ticker Setup ---
ticker_symbols = ['PEP', 'KO']
ticker = st.selectbox("Select Stock", ticker_symbols)
model_type = st.radio("Model Type", ["LSTM", "XGBoost"], horizontal=True)

# --- Feature Input (XGBoost only) ---
xgb_features = ['Adj Close', 'Volume', 'RSI', 'SMA_20', 'MACD']
xgb_inputs = {}

if model_type == "XGBoost":
    st.subheader("📥 Enter Feature Values for XGBoost Prediction")
    for feat in xgb_features:
        xgb_inputs[feat] = st.number_input(f"{feat}:", value=50.0, step=0.1)

# --- Prediction Button ---
if st.button("🔮 Predict Next 30 Days" if model_type == "LSTM" else "🔢 Predict Now"):
    try:
        df = yf.download(ticker, period="5y", interval="1d")
        df = df[["Close"]].dropna()

        if model_type == "LSTM":
            if lstm_models[ticker] is None:
                st.error(f"⚠️ LSTM model for {ticker} not found.")
            else:
                scaler = MinMaxScaler()
                scaled = scaler.fit_transform(df)
                last_60 = scaled[-60:].reshape(1, 60, 1)
                model = lstm_models[ticker]

                predictions = []
                for _ in range(30):
                    pred = model.predict(last_60, verbose=0)[0][0]
                    predictions.append(pred)
                    last_60 = np.append(last_60[:, 1:, :], [[[pred]]], axis=1)

                forecasted = scaler.inverse_transform(np.array(predictions).reshape(-1, 1))
                forecast_df = pd.DataFrame(forecasted, columns=["Forecasted Price"])
                st.success(f"✅ 30-Day Forecast for {ticker} Complete")
                st.line_chart(forecast_df)
                st.session_state["results"] = {
                    "days": np.arange(1, 31),
                    "primary": forecast_df["Forecasted Price"].values,
                    "primary_stock": ticker,
                    "competitor_stock": "N/A",
                    "advice": "Consider Buying 🚀" if forecast_df.iloc[-1, 0] > forecast_df.iloc[0, 0] else "Hold / Watch ⏸️"
                }

        elif model_type == "XGBoost":
            if xgb_models[ticker] is None:
                st.error(f"⚠️ XGBoost model for {ticker} not found.")
            else:
                x_input = np.array([list(xgb_inputs.values())])
                dmatrix = xgb.DMatrix(x_input, feature_names=xgb_features)
                prediction = xgb_models[ticker].predict(dmatrix)
                st.success(f"📈 {ticker} Forecasted Value: ${prediction[0]:.2f}")
                st.session_state["results"] = {
                    "days": [1],
                    "primary": prediction,
                    "primary_stock": ticker,
                    "competitor_stock": "N/A",
                    "advice": "XGBoost single-day forecast"
                }

    except Exception as e:
        st.error(f"An error occurred during forecasting: {e}")

# --- Tabs for Analysis and Download ---
tab1, tab2 = st.tabs(["📊 Analysis", "📥 Download Report"])

with tab1:
    result = st.session_state.get("results", {})
    if result:
        df = pd.DataFrame({
            "Day": result["days"],
            result["primary_stock"]: np.ravel(result["primary"])
        })
        fig = px.line(df, x="Day", y=df.columns[1], title="Forecasted Price Trend", markers=True)
        st.plotly_chart(fig)
        st.info(f"💡 Advisory: {result['advice']}")
    else:
        st.warning("🔎 No predictions yet. Run a forecast above.")

with tab2:
    result = st.session_state.get("results", {})
    if result:
        df = pd.DataFrame({
            "Day": result["days"],
            result["primary_stock"]: np.ravel(result["primary"])
        })

        # Excel
        df.to_excel("forecast.xlsx", index=False)
        with open("forecast.xlsx", "rb") as f:
            st.download_button("📊 Download Excel", f, file_name="forecast.xlsx")

        # PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Forecast Report", ln=True, align="C")
        for i, val in enumerate(result["primary"]):
            pdf.cell(200, 10, txt=f"Day {i+1}: {result['primary_stock']} = {val:.2f}".encode('latin-1', 'replace').decode('latin-1'), ln=True)
        pdf.output("forecast.pdf")
        with open("forecast.pdf", "rb") as f:
            st.download_button("📄 Download PDF", f, file_name="forecast.pdf")

        # Word DOC
        doc = Document()
        doc.add_heading("Stock Forecast Report", 0)
        doc.add_paragraph(f"Forecast for {result['primary_stock']}")
        for i, val in enumerate(result["primary"]):
            doc.add_paragraph(f"Day {i+1}: {val:.2f}")
        doc.add_paragraph(f"Advisory: {result['advice']}")
        doc.save("forecast.docx")
        with open("forecast.docx", "rb") as f:
            st.download_button("📝 Download DOC", f, file_name="forecast.docx")
    else:
        st.warning("📁 Generate prediction first to enable downloads.")
