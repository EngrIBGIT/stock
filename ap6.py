import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import matplotlib.pyplot as plt
import requests
import time
from fpdf import FPDF
from docx import Document
import yfinance as yf
import os
from tensorflow.keras.models import load_model
import xgboost as xgb
from sklearn.preprocessing import MinMaxScaler
import talib

# --- Set Page Config (Must be first Streamlit call) ---
st.set_page_config(page_title="PEP vs KO Stock Prediction", layout="wide")

# --- Theme Colors ---
st.markdown("""
<style>
.main {background-color: #e0f7fa;}
.sidebar .sidebar-content {background-color: #d8f3dc;}
.css-1q8dd3e {background-color: #fff3cd;}
.marquee {
    font-size: 18px;
    color: #0c0c0c;
    padding: 10px;
    background-color: #ffffe0;
    overflow: hidden;
    white-space: nowrap;
    animation: marquee 20s linear infinite;
}
@keyframes marquee {
    0% {transform: translateX(100%);}
    100% {transform: translateX(-100%);}
}
</style>
""", unsafe_allow_html=True)

# --- API Keys ---
FINNHUB_API_KEY = "cvuguhpr01qjg139orhgcvuguhpr01qjg139ori0"
MARKETSTACK_API_KEY = "359c056969eeb01527ce644a4df15822"

# --- Load Models Automatically ---
@st.cache_resource
def load_lstm_model(path):
    return load_model(path, compile=False)

@st.cache_resource
def load_xgb_model(path):
    model = xgb.Booster()
    model.load_model(path)
    return model

# Load models from local files with proper error handling
try:
    lstm_models = {
        "PEP": load_lstm_model("models/pep_lstm_model.h5") if os.path.exists("models/pep_lstm_model.h5") else None,
        "KO": load_lstm_model("models/ko_lstm_model.h5") if os.path.exists("models/ko_lstm_model.h5") else None
    }

    xgb_models = {
        "PEP": load_xgb_model("models/xgb_pep_model.xgb") if os.path.exists("models/xgb_pep_model.xgb") else None,
        "KO": load_xgb_model("models/xgb_ko_model.xgb") if os.path.exists("models/xgb_ko_model.xgb") else None
    }
except Exception as e:
    st.error(f"Error loading models: {e}")
    lstm_models = {}
    xgb_models = {}

# --- Feature Engineering Functions ---
def calculate_technical_indicators(df):
    """Calculate all required technical indicators for the models"""
    df['SMA_20'] = talib.SMA(df['Close'], timeperiod=20)
    df['EMA_20'] = talib.EMA(df['Close'], timeperiod=20)
    df['RSI'] = talib.RSI(df['Close'], timeperiod=14)
    macd, signal, _ = talib.MACD(df['Close'])
    df['MACD'] = macd
    df['Signal_Line'] = signal
    df['Daily_Return'] = df['Close'].pct_change()
    df['Volatility'] = df['Daily_Return'].rolling(window=20).std()
    upper, middle, lower = talib.BBANDS(df['Close'], timeperiod=20)
    df['Middle_Band'] = middle
    df['Std_Dev'] = (upper - lower) / 2
    return df.dropna()

# --- Cached API with retry ---
@st.cache_data(ttl=600)
def get_stock_price(symbol):
    try:
        url = f"https://finnhub.io/api/v1/quote?symbol={symbol}&token={FINNHUB_API_KEY}"
        for _ in range(3):
            response = requests.get(url, timeout=5)
            if response.status_code == 200:
                return response.json().get('c', 'N/A')
            time.sleep(1)
    except Exception:
        return "Error"
    return "Error"

@st.cache_data(ttl=600)
def get_financial_news():
    try:
        url = f"http://api.marketstack.com/v1/news?access_key={MARKETSTACK_API_KEY}"
        response = requests.get(url)
        if response.status_code == 200:
            return response.json().get('data', [])
    except Exception:
        return []
    return []

# --- Setup Stock Ticker ---
st.title("📈 PEP vs KO Stock Prediction App")

# Only PEP and KO stocks
primary_tickers = ["PEP"]
competitor_tickers = ["KO"]

primary_prices = [get_stock_price(sym) for sym in primary_tickers]
competitor_prices = [get_stock_price(sym) for sym in competitor_tickers]

primary_ticker_str = " | ".join([f"{sym}: ${price}" for sym, price in zip(primary_tickers, primary_prices)])
competitor_ticker_str = " | ".join([f"{sym}: ${price}" for sym, price in zip(competitor_tickers, competitor_prices)])

st.markdown(f'<div class="marquee">🔔 Primary Stock: {primary_ticker_str}</div>', unsafe_allow_html=True)
st.markdown(f'<div class="marquee">🔔 Competitor Stock: {competitor_ticker_str}</div>', unsafe_allow_html=True)

# --- Tabs ---
tab1, tab2, tab3, tab4, tab5 = st.tabs(["Home", "Prediction", "Feature-Based Prediction", "Analysis", "Download Reports"])

# --- Home Tab ---
with tab1:
    st.header("📘 Welcome to the PEP vs KO Stock Prediction App")
    st.markdown("""
    🌟 This app enables stock price predictions and comparative analysis between PEP and KO stocks.
    Gain insights into market trends, visualize stock comparisons, and receive advisory on selected stocks. 
    
    🌟 Select models, input variables, and enjoy visualizations
    of market trends!

    Features:
    - 🔄 Real-time stock prices
    - 🧠 ML-powered prediction interface
    - 📊 Interactive analysis and advisory
    - 📰 Financial news
    - 📥 Download Excel, PDF, DOC reports

    This app lets you analyze and predict stock trends using dynamic price estimations from yfinance (Yahoo Finance).
    """)

    st.subheader("📰 Latest News")
    news = get_financial_news()
    if news:
        for article in news[:5]:
            st.markdown(f"**{article['title']}**  \n{article['description']}  \n[Read more]({article['url']})")
            st.markdown("---")
    else:
        st.warning("News unavailable.")

# --- Prediction Tab ---
with tab2:
    st.header("📊 Quick Prediction")

    col1, col2 = st.columns(2)
    with col1:
        primary_stock = st.selectbox("Primary Stock", primary_tickers)
    with col2:
        competitor_stock = st.selectbox("Competitor Stock", competitor_tickers)
    
    model_choice = st.radio("Model", ["XGBoost", "LSTM"])

    # --- Button to trigger Prediction --- 
    if st.button("🔮 Predict Next 30 Days"):
        try:
            # Get data for both stocks
            df_primary = yf.download(primary_stock, period="5y", interval="1d")
            df_competitor = yf.download(competitor_stock, period="5y", interval="1d")
            
            if df_primary.empty or df_competitor.empty:
                st.error("Could not download stock data. Please try again later.")
                st.stop()
            
            # Calculate technical indicators
            df_primary = calculate_technical_indicators(df_primary)
            df_competitor = calculate_technical_indicators(df_competitor)
            
            # Initialize forecast variables
            forecasted_primary = None
            forecasted_competitor = None
            
            if model_choice == "LSTM":
                # Process primary stock
                if primary_stock in lstm_models and lstm_models[primary_stock] is not None:
                    scaler_primary = MinMaxScaler()
                    scaled_primary = scaler_primary.fit_transform(df_primary[['Close']])
                    last_60_primary = scaled_primary[-60:].reshape(1, 60, 1)
                    
                    model_primary = lstm_models[primary_stock]
                    predictions_primary = []
                    for _ in range(30):
                        pred = model_primary.predict(last_60_primary, verbose=0)[0][0]
                        predictions_primary.append(pred)
                        last_60_primary = np.append(last_60_primary[:, 1:, :], [[[pred]]], axis=1)
                    forecasted_primary = scaler_primary.inverse_transform(np.array(predictions_primary).reshape(-1, 1))
                else:
                    st.error(f"⚠️ LSTM model for {primary_stock} not found. Using random predictions.")
                    forecasted_primary = np.random.uniform(df_primary['Close'].min(), df_primary['Close'].max(), 30).reshape(-1, 1)
                
                # Process competitor stock
                if competitor_stock in lstm_models and lstm_models[competitor_stock] is not None:
                    scaler_competitor = MinMaxScaler()
                    scaled_competitor = scaler_competitor.fit_transform(df_competitor[['Close']])
                    last_60_competitor = scaled_competitor[-60:].reshape(1, 60, 1)
                    
                    model_competitor = lstm_models[competitor_stock]
                    predictions_competitor = []
                    for _ in range(30):
                        pred = model_competitor.predict(last_60_competitor, verbose=0)[0][0]
                        predictions_competitor.append(pred)
                        last_60_competitor = np.append(last_60_competitor[:, 1:, :], [[[pred]]], axis=1)
                    forecasted_competitor = scaler_competitor.inverse_transform(np.array(predictions_competitor).reshape(-1, 1))
                else:
                    st.error(f"⚠️ LSTM model for {competitor_stock} not found. Using random predictions.")
                    forecasted_competitor = np.random.uniform(df_competitor['Close'].min(), df_competitor['Close'].max(), 30).reshape(-1, 1)

            elif model_choice == "XGBoost":
                # Define required features for XGBoost models
                required_features = ['High', 'Volume', 'Adj Close', 'Daily_Return', 'SMA_20', 
                                   'EMA_20', 'RSI', 'MACD', 'Signal_Line', 'Volatility', 
                                   'Middle_Band', 'Std_Dev']
                
                # Process primary stock
                if primary_stock in xgb_models and xgb_models[primary_stock] is not None:
                    try:
                        # Prepare input data with all required features
                        input_data_primary = df_primary[required_features].iloc[-1:].values
                        dmatrix_primary = xgb.DMatrix(input_data_primary, feature_names=required_features)
                        prediction_primary = xgb_models[primary_stock].predict(dmatrix_primary)
                        forecasted_primary = np.full(30, prediction_primary[0]).reshape(-1, 1)
                    except Exception as e:
                        st.error(f"Error predicting {primary_stock}: {str(e)}")
                        forecasted_primary = np.random.uniform(df_primary['Close'].min(), df_primary['Close'].max(), 30).reshape(-1, 1)
                
                # Process competitor stock
                if competitor_stock in xgb_models and xgb_models[competitor_stock] is not None:
                    try:
                        input_data_competitor = df_competitor[required_features].iloc[-1:].values
                        dmatrix_competitor = xgb.DMatrix(input_data_competitor, feature_names=required_features)
                        prediction_competitor = xgb_models[competitor_stock].predict(dmatrix_competitor)
                        forecasted_competitor = np.full(30, prediction_competitor[0]).reshape(-1, 1)
                    except Exception as e:
                        st.error(f"Error predicting {competitor_stock}: {str(e)}")
                        forecasted_competitor = np.random.uniform(df_competitor['Close'].min(), df_competitor['Close'].max(), 30).reshape(-1, 1)

            # Store results
            st.session_state["results"] = {
                "primary": forecasted_primary.flatten(),
                "competitor": forecasted_competitor.flatten(),
                "primary_stock": primary_stock,
                "competitor_stock": competitor_stock,
                "advice": "Consider Buying 🚀" if forecasted_primary[-1][0] > forecasted_primary[0][0] else "Hold / Watch ⏸️"
            }

            # Display results
            col1, col2 = st.columns(2)
            with col1:
                st.success(f"📈 {primary_stock} Forecast:")
                st.line_chart(pd.DataFrame(forecasted_primary, columns=["Forecasted Price"]))
            with col2:
                st.success(f"📉 {competitor_stock} Forecast:")
                st.line_chart(pd.DataFrame(forecasted_competitor, columns=["Forecasted Price"]))
            
            st.info(f"💡 Advisory: {st.session_state['results']['advice']}")

        except Exception as e:
            st.error(f"An error occurred during forecasting: {str(e)}")

# --- Feature-Based Prediction ---
with tab3:
    st.header("🔍 Predict by Feature Set")
    model_type = st.selectbox("Model Type", ["XGBoost", "LSTM"])

    # Only show feature inputs for XGBoost
    if model_type == "XGBoost":
        feature_set = ['High', 'Volume', 'Adj Close', 'Daily_Return', 'SMA_20', 
                      'EMA_20', 'RSI', 'MACD', 'Signal_Line', 'Volatility', 
                      'Middle_Band', 'Std_Dev']
        feature_inputs = {}
        for feat in feature_set:
            feature_inputs[feat] = st.number_input(f"{feat}:", value=0.0, step=0.1, key=f"xgb_{feat}")
    else:
        st.info("For LSTM predictions, please use the 'Quick Prediction' tab.")

    if st.button("🔢 Predict Based on Features") and model_type == "XGBoost":
        try:
            # Prepare input data
            input_data = np.array([list(feature_inputs.values())])
            dmatrix = xgb.DMatrix(input_data, feature_names=list(feature_inputs.keys()))
            
            # Get predictions for all available models
            predictions = {}
            for stock, model in xgb_models.items():
                if model is not None:
                    predictions[stock] = model.predict(dmatrix)[0]
            
            if predictions:
                st.success("Predictions:")
                for stock, pred in predictions.items():
                    st.write(f"{stock}: ${pred:.2f}")
                
                # Store results for visualization
                st.session_state["results"] = {
                    "primary": np.full(30, list(predictions.values())[0]),
                    "competitor": np.full(30, list(predictions.values())[-1]),
                    "primary_stock": list(predictions.keys())[0],
                    "competitor_stock": list(predictions.keys())[-1],
                    "advice": "Consider Buying 🚀" if list(predictions.values())[0] > list(predictions.values())[-1] else "Hold / Watch ⏸️"
                }
            else:
                st.warning("No XGBoost models available for prediction.")
        except Exception as e:
            st.error(f"Error during prediction: {str(e)}")

# --- Analysis Tab ---
with tab4:
    st.header("📊 Comparative Analysis")
    result = st.session_state.get("results", {})
    if result:
        primary = np.ravel(result['primary'])
        competitor = np.ravel(result['competitor'])
        df = pd.DataFrame({
            "Day": np.arange(1, len(primary) + 1),
            result['primary_stock']: primary,
            result['competitor_stock']: competitor
        })
        df_melted = df.melt(id_vars="Day", var_name="Stock", value_name="Price")
        fig = px.line(df_melted, x="Day", y="Price", color="Stock", markers=True)
        st.plotly_chart(fig)
        st.success(f"🧠 Advisory: {result['advice']}")
    else:
        st.warning("Run predictions to see analysis.")

# --- Download Reports Tab ---
with tab5:
    st.header("📥 Download Reports")
    result = st.session_state.get("results", {})
    if result:
        primary = np.ravel(result['primary'])
        competitor = np.ravel(result['competitor'])

        df = pd.DataFrame({
            "Day": np.arange(1, len(primary) + 1),
            result['primary_stock']: primary,
            result['competitor_stock']: competitor
        })

        # Excel
        df.to_excel("report.xlsx", index=False)
        with open("report.xlsx", "rb") as f:
            st.download_button("📊 Download Excel", f, file_name="stock_analysis.xlsx")

        # PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Stock Analysis Report", ln=True, align="C")
        for i in range(len(primary)):
            line = f"Day {i+1}: {result['primary_stock']}={primary[i]:.2f} | {result['competitor_stock']}={competitor[i]:.2f}"
            pdf.cell(200, 10, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
        pdf.cell(200, 10, txt=f"Advisory: {result['advice']}".encode('latin-1', 'replace').decode('latin-1'), ln=True)
        pdf.output("report.pdf")
        with open("report.pdf", "rb") as f:
            st.download_button("📄 Download PDF", f, file_name="stock_report.pdf")

        # DOC
        doc = Document()
        doc.add_heading("Stock Analysis Report", 0)
        doc.add_paragraph(f"Advisory: {result['advice']}")
        for i in range(len(primary)):
            doc.add_paragraph(f"Day {i+1}: {result['primary_stock']}={primary[i]:.2f}, {result['competitor_stock']}={competitor[i]:.2f}")
        doc.save("report.docx")
        with open("report.docx", "rb") as f:
            st.download_button("📄 Download DOC", f, file_name="stock_report.docx")
    else:
        st.warning("Generate predictions to enable downloads.")

# --- Footer ---
st.markdown("""
**Note:** Predictions are based on historical price data using Yahoo Finance and are for informational purposes only.
""")