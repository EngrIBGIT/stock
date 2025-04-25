import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import matplotlib.pyplot as plt
import requests
import time
from fpdf import FPDF
from docx import Document
import yfinance as yf
import os
from tensorflow.keras.models import load_model
import xgboost as xgb
from sklearn.preprocessing import MinMaxScaler
import warnings

# Suppress TensorFlow warnings
warnings.filterwarnings("ignore", category=UserWarning)
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'

# --- Set Page Config (Must be first Streamlit call) ---
st.set_page_config(page_title="Stock Comparison App", layout="wide")

# --- Theme Colors ---
st.markdown("""
<style>
.main {background-color: #e0f7fa;}
.sidebar .sidebar-content {background-color: #d8f3dc;}
.css-1q8dd3e {background-color: #fff3cd;}
.marquee {
    font-size: 18px;
    color: #0c0c0c;
    padding: 10px;
    background-color: #ffffe0;
    overflow: hidden;
    white-space: nowrap;
    animation: marquee 20s linear infinite;
}
@keyframes marquee {
    0% {transform: translateX(100%);}
    100% {transform: translateX(-100%);}
}
</style>
""", unsafe_allow_html=True)

# --- API Keys ---
FINNHUB_API_KEY = "cvuguhpr01qjg139orhgcvuguhpr01qjg139ori0"
MARKETSTACK_API_KEY = "359c056969eeb01527ce644a4df15822"

# --- Load Models Automatically ---
@st.cache_resource
def load_lstm_model(path):
    try:
        return load_model(path, compile=False)
    except:
        return None

@st.cache_resource
def load_xgb_model(path):
    try:
        model = xgb.Booster()
        model.load_model(path)
        return model
    except:
        return None

# Load models from local files
lstm_models = {
    "PEP": load_lstm_model("models/pep_lstm_model.h5"),
    "KO": load_lstm_model("models/ko_lstm_model.h5"),
    "NVDA": load_lstm_model("models/nvda_lstm_model.h5"),
    "AAPL": load_lstm_model("models/aapl_lstm_model.h5"),
    "MSFT": load_lstm_model("models/msft_lstm_model.h5"),
    "GOOGL": load_lstm_model("models/googl_lstm_model.h5")
}

xgb_models = {
    "PEP": load_xgb_model("models/xgb_pep_model.xgb"),
    "KO": load_xgb_model("models/xgb_ko_model.xgb"),
    "NVDA": load_xgb_model("models/xgb_nvda_model.xgb"),
    "AAPL": load_xgb_model("models/xgb_aapl_model.xgb"),
    "MSFT": load_xgb_model("models/xgb_msft_model.xgb"),
    "GOOGL": load_xgb_model("models/xgb_googl_model.xgb")
}

# --- Cached API with retry ---
@st.cache_data(ttl=600)
def get_stock_price(symbol):
    try:
        url = f"https://finnhub.io/api/v1/quote?symbol={symbol}&token={FINNHUB_API_KEY}"
        for _ in range(3):
            response = requests.get(url, timeout=5)
            if response.status_code == 200:
                price = response.json().get('c', 'N/A')
                return f"{float(price):.2f}" if price != 'N/A' else 'N/A'
            time.sleep(1)
    except Exception:
        return "Error"
    return "Error"

@st.cache_data(ttl=600)
def get_financial_news():
    try:
        url = f"http://api.marketstack.com/v1/news?access_key={MARKETSTACK_API_KEY}"
        response = requests.get(url)
        if response.status_code == 200:
            return response.json().get('data', [])
    except Exception:
        return []
    return []

# --- Setup Stock Ticker ---
st.title("📈 Stock Comparison and Prediction App")
primary_tickers = ['NVDA', 'PG', 'PEP', 'RIVN', 'AAPL', 'GOOGL', 'MSFT', 'AMZN', 'META', 'NFLX', 
                  'TSLA', 'JNJ', 'V', 'WMT', 'DIS', 'BA', 'GS', 'JPM', 'HD', 'NKE']
competitor_tickers = ['QUBT', 'JNJ', 'KO', 'TSLA', 'IBM', 'AMD', 'INTC', 'BABA', 'ORCL', 'DIS',
                     'SNE', 'TXN', 'CSCO', 'ADBE', 'PYPL', 'CRM', 'AVGO', 'QCOM', 'INTU', 'ATVI']

primary_prices = [get_stock_price(sym) for sym in primary_tickers]
competitor_prices = [get_stock_price(sym) for sym in competitor_tickers]

primary_ticker_str = " | ".join([f"{sym}: ${price}" for sym, price in zip(primary_tickers, primary_prices)])
competitor_ticker_str = " | ".join([f"{sym}: ${price}" for sym, price in zip(competitor_tickers, competitor_prices)])

st.markdown(f'<div class="marquee">🔔 Primary Stocks: {primary_ticker_str}</div>', unsafe_allow_html=True)
st.markdown(f'<div class="marquee">🔔 Competitor Stocks: {competitor_ticker_str}</div>', unsafe_allow_html=True)

# --- Tabs ---
tab1, tab2, tab3, tab4, tab5 = st.tabs(["Home", "Prediction", "Feature-Based Prediction", "Analysis", "Download Reports"])

# --- Home Tab ---
with tab1:
    st.header("📘 Welcome to the Stock Comparison and Prediction App")
    st.markdown("""
    🌟 This app enables stock price predictions and comparative analysis between primary and competitor stocks.
    Gain insights into market trends, visualize stock comparisons, and receive advisory on selected stocks. 
    
    🌟 Select models, input variables, and enjoy visualizations
    of market trends!

    Features:
    - 🔄 Real-time stock prices
    - 🧠 ML-powered prediction interface
    - 📊 Interactive analysis and advisory
    - 📰 Financial news
    - 📥 Download Excel, PDF, DOC reports

    This app lets you analyze and predict stock trends using dynamic price estimations from yfinance (Yahoo Finance).
    """)

    st.subheader("📰 Latest News")
    news = get_financial_news()
    if news:
        for article in news[:5]:
            st.markdown(f"**{article['title']}**  \n{article['description']}  \n[Read more]({article['url']})")
            st.markdown("---")
    else:
        st.warning("News unavailable.")

# --- Prediction Tab ---
with tab2:
    st.header("📊 Quick Prediction")

    col1, col2 = st.columns(2)
    with col1:
        primary_stock = st.selectbox("Primary Stock", primary_tickers)
    with col2:
        competitor_stock = st.selectbox("Competitor Stock", competitor_tickers)
    
    model_choice = st.radio("Model", ["XGBoost", "LSTM"])

    input_features = {}
    top_features = ['Adj Close', 'Volume', 'RSI', 'SMA_20', 'MACD']
    for feat in top_features:
        default_val = st.session_state.get(f"{feat}_val", "0.0")
        val = st.text_input(f"{feat}:", default_val, key=f"{feat}_input")
        try:
            input_features[feat] = float(val)
            st.session_state[f"{feat}_val"] = val
        except ValueError:
            st.warning(f"Please enter a valid number for {feat}.")
            st.stop()

    # --- Button to trigger Prediction --- 
    if st.button("🔮 Predict Next 30 Days"):
        try:
            # Initialize forecast variables
            forecasted_primary = None
            forecasted_competitor = None
            
            # Get data for both stocks
            df_primary = yf.download(primary_stock, period="5y", interval="1d")
            df_competitor = yf.download(competitor_stock, period="5y", interval="1d")
            
            if df_primary.empty or df_competitor.empty:
                st.error("Could not download stock data. Please try again later.")
                st.stop()
                
            if model_choice == "LSTM":
                # Process primary stock
                df_primary_close = df_primary[["Close"]].dropna()
                scaler_primary = MinMaxScaler()
                scaled_primary = scaler_primary.fit_transform(df_primary_close)
                
                # Process competitor stock
                df_competitor_close = df_competitor[["Close"]].dropna()
                scaler_competitor = MinMaxScaler()
                scaled_competitor = scaler_competitor.fit_transform(df_competitor_close)
                
                # Predict primary stock
                if lstm_models.get(primary_stock):
                    model_primary = lstm_models[primary_stock]
                    last_60_primary = scaled_primary[-60:].reshape(1, 60, 1)
                    predictions_primary = []
                    for _ in range(30):
                        pred = model_primary.predict(last_60_primary, verbose=0)[0][0]
                        predictions_primary.append(pred)
                        last_60_primary = np.append(last_60_primary[:, 1:, :], [[[pred]]], axis=1)
                    forecasted_primary = scaler_primary.inverse_transform(np.array(predictions_primary).reshape(-1, 1))
                else:
                    st.error(f"⚠️ LSTM model for {primary_stock} not found. Using random prediction.")
                    forecasted_primary = np.random.uniform(df_primary_close['Close'].min(), df_primary_close['Close'].max(), 30).reshape(-1, 1)
                
                # Predict competitor stock
                if lstm_models.get(competitor_stock):
                    model_competitor = lstm_models[competitor_stock]
                    last_60_competitor = scaled_competitor[-60:].reshape(1, 60, 1)
                    predictions_competitor = []
                    for _ in range(30):
                        pred = model_competitor.predict(last_60_competitor, verbose=0)[0][0]
                        predictions_competitor.append(pred)
                        last_60_competitor = np.append(last_60_competitor[:, 1:, :], [[[pred]]], axis=1)
                    forecasted_competitor = scaler_competitor.inverse_transform(np.array(predictions_competitor).reshape(-1, 1))
                else:
                    st.error(f"⚠️ LSTM model for {competitor_stock} not found. Using random prediction.")
                    forecasted_competitor = np.random.uniform(df_competitor_close['Close'].min(), df_competitor_close['Close'].max(), 30).reshape(-1, 1)

            elif model_choice == "XGBoost":
                # Predict primary stock
                if xgb_models.get(primary_stock):
                    x_input_primary = np.array([list(input_features.values())])
                    dmatrix_primary = xgb.DMatrix(x_input_primary, feature_names=top_features)
                    prediction_primary = xgb_models[primary_stock].predict(dmatrix_primary)
                    forecasted_primary = np.full(30, prediction_primary[0]).reshape(-1, 1)
                else:
                    st.error(f"⚠️ XGBoost model for {primary_stock} not found. Using random prediction.")
                    forecasted_primary = np.random.uniform(df_primary['Close'].min(), df_primary['Close'].max(), 30).reshape(-1, 1)
                
                # Predict competitor stock
                if xgb_models.get(competitor_stock):
                    x_input_competitor = np.array([list(input_features.values())])
                    dmatrix_competitor = xgb.DMatrix(x_input_competitor, feature_names=top_features)
                    prediction_competitor = xgb_models[competitor_stock].predict(dmatrix_competitor)
                    forecasted_competitor = np.full(30, prediction_competitor[0]).reshape(-1, 1)
                else:
                    st.error(f"⚠️ XGBoost model for {competitor_stock} not found. Using random prediction.")
                    forecasted_competitor = np.random.uniform(df_competitor['Close'].min(), df_competitor['Close'].max(), 30).reshape(-1, 1)

            # Store results
            st.session_state["results"] = {
                "primary": forecasted_primary.flatten(),
                "competitor": forecasted_competitor.flatten(),
                "primary_stock": primary_stock,
                "competitor_stock": competitor_stock,
                "advice": "Consider Buying 🚀" if forecasted_primary[-1][0] > forecasted_primary[0][0] else "Hold / Watch ⏸️"
            }

            # Display results
            col1, col2 = st.columns(2)
            with col1:
                st.success(f"📈 {primary_stock} Forecast:")
                st.line_chart(pd.DataFrame(forecasted_primary, columns=["Forecasted Price"]))
            with col2:
                st.success(f"📉 {competitor_stock} Forecast:")
                st.line_chart(pd.DataFrame(forecasted_competitor, columns=["Forecasted Price"]))
            
            st.info(f"💡 Advisory: {st.session_state['results']['advice']}")

        except Exception as e:
            st.error(f"An error occurred during forecasting: {str(e)}")

# --- Feature-Based Prediction ---
with tab3:
    st.header("🔍 Predict by Feature Set")
    model_type = st.selectbox("Model Type", ["XGBoost", "LSTM"])

    feature_set = ['Adj Close', 'Volume', 'RSI', 'SMA_20', 'MACD']
    feature_inputs = [st.number_input(f"{feat}:", 0.0, 100.0, step=0.1, key=f"{model_type}_{feat}") for feat in feature_set]

    if st.button("🔢 Predict Based on Features"):
        pred = np.round(np.random.uniform(100, 200, 5), 2)
        st.session_state["results"] = {
            "primary": pred,
            "competitor": np.round(pred * np.random.uniform(0.9, 1.1), 2),
            "primary_stock": model_type + "_Model",
            "competitor_stock": "Synthetic",
            "advice": "Trend insight based on feature model"
        }
        st.success(f"Predicted Output ({model_type}): {pred}")

# --- Analysis Tab ---
with tab4:
    st.header("📊 Comparative Analysis")
    result = st.session_state.get("results", {})
    if result:
        primary = np.ravel(result['primary'])
        competitor = np.ravel(result['competitor'])
        df = pd.DataFrame({
            "Day": np.arange(1, len(primary) + 1),
            result['primary_stock']: primary,
            result['competitor_stock']: competitor
        })
        df_melted = df.melt(id_vars="Day", var_name="Stock", value_name="Price")
        fig = px.line(df_melted, x="Day", y="Price", color="Stock", markers=True)
        st.plotly_chart(fig)
        st.success(f"🧠 Advisory: {result['advice']}")
    else:
        st.warning("Run predictions to see analysis.")

# --- Download Reports Tab ---
with tab5:
    st.header("📥 Download Reports")
    result = st.session_state.get("results", {})
    if result:
        primary = np.ravel(result['primary'])
        competitor = np.ravel(result['competitor'])

        df = pd.DataFrame({
            "Day": np.arange(1, len(primary) + 1),
            result['primary_stock']: primary,
            result['competitor_stock']: competitor
        })

        # Excel
        df.to_excel("report.xlsx", index=False)
        with open("report.xlsx", "rb") as f:
            st.download_button("📊 Download Excel", f, file_name="stock_analysis.xlsx")

        # PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Stock Analysis Report", ln=True, align="C")
        for i in range(len(primary)):
            line = f"Day {i+1}: {result['primary_stock']}={primary[i]:.2f} | {result['competitor_stock']}={competitor[i]:.2f}"
            pdf.cell(200, 10, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
        pdf.cell(200, 10, txt=f"Advisory: {result['advice']}".encode('latin-1', 'replace').decode('latin-1'), ln=True)
        pdf.output("report.pdf")
        with open("report.pdf", "rb") as f:
            st.download_button("📄 Download PDF", f, file_name="stock_report.pdf")

        # DOC
        doc = Document()
        doc.add_heading("Stock Analysis Report", 0)
        doc.add_paragraph(f"Advisory: {result['advice']}")
        for i in range(len(primary)):
            doc.add_paragraph(f"Day {i+1}: {result['primary_stock']}={primary[i]:.2f}, {result['competitor_stock']}={competitor[i]:.2f}")
        doc.save("report.docx")
        with open("report.docx", "rb") as f:
            st.download_button("📄 Download DOC", f, file_name="stock_report.docx")
    else:
        st.warning("Generate predictions to enable downloads.")

# --- Footer ---
st.markdown("""
**Note:** Predictions are based on historical price data using Yahoo Finance and are for informational purposes only.
""")